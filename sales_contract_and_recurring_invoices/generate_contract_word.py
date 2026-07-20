"""
HHS Contract Document - Word Generator
Generates a bilingual EN/AR Word DOCX directly from contract JSON data.
Usage: python generate_contract_word.py <data.json> <output.docx>
"""
import sys, os, json
sys.path.insert(0, r'C:\Odoo_Libraries')

try:
    import docx
    print(f"DEBUG DOCX LOC: {getattr(docx, '__file__', 'unknown')}", file=sys.stderr)
    print(f"DEBUG SYS.PATH: {sys.path}", file=sys.stderr)
    from docx import Document
except Exception as e:
    import docx
    print(f"DEBUG FATAL: {e}. DOCX path: {getattr(docx, '__path__', 'no-path')}", file=sys.stderr)
    raise
from docx.shared import Inches, Pt, RGBColor, Cm, Twips, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

if len(sys.argv) < 3:
    print("Usage: python generate_contract_word.py <data.json> <output.docx>")
    sys.exit(1)

json_file, output_file = sys.argv[1], sys.argv[2]
with open(json_file, 'r', encoding='utf-8') as f:
    d = json.load(f)

# ── Constants ─────────────────────────────────────────────────────────────
NAVY       = RGBColor(0x00, 0x20, 0x60)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX   = '002060'
HDR_BG_HEX = '1F4E78'
IMG_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static', 'src', 'img')
LOGO_PATH        = os.path.join(IMG_DIR, 'HHS_Header_Middle.png')
COVER_IMG_PATH   = os.path.join(IMG_DIR, 'image_cn.png')
FOOTER1_IMG_PATH = os.path.join(IMG_DIR, 'ContractDocument_1stPage_Footer.png')

# ── Data ──────────────────────────────────────────────────────────────────
today_str             = d.get('today', '')
ref_no                = d.get('name', '')
partner_name          = d.get('partner_name', '')
customer_name         = d.get('customer_name', '')
id_party              = d.get('id_party', '')
job_position          = d.get('job_position', '')
mobile_no             = d.get('mobile_no', '')
email_addr            = d.get('email', '')
date_start            = d.get('date_start', '')
date_end              = d.get('date_end', '')
amount_total          = d.get('amount_total', 0)
payment_term_text     = d.get('payment_term_text', '')
payment_term_text_ar  = d.get('payment_term_text_ar', '')
payment_schedule_lines= d.get('payment_schedule_lines', [])
add_paid_price        = d.get('add_paid_service_price', 0)
contract_lines        = d.get('contract_lines', [])
svc_coord_person      = d.get('service_coordinator_person', '')
svc_coord_mobile      = d.get('service_coordinator_mobile', '')
contact_persons       = d.get('contact_persons', '')
contact_persons_mobile= d.get('contact_persons_mobile', '')
additional_info       = d.get('additional_info', '')
site_address          = d.get('site_address', '')
total_emergency       = int(sum(l.get('no_of_emergency_visit', 0) or 0 for l in contract_lines))

# Deduplicated visits for section 2.1
seen_visits = {}
for l in contract_lines:
    key = (l.get('description',''), l.get('no_of_visits_per_year', 0))
    if key not in seen_visits:
        seen_visits[key] = l

# ── Helpers ───────────────────────────────────────────────────────────────
def set_cell_margins(cell, top=None, bottom=None, left=None, right=None):
    """Set cell margins (padding) in twentieths of a point (dxa) for a cell."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for name, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val is not None:
            node = OxmlElement(f'w:{name}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
    tcPr.append(tcMar)

def sort_pPr(pPr):
    """Sort elements in paragraph properties pPr strictly matching the Word OOXML schema to prevent validation errors."""
    order = [
        'pStyle', 'keepNext', 'keepLines', 'pageBreakBefore', 'framePr',
        'widowControl', 'numPr', 'pBdr', 'shd', 'tabs', 'spacing', 'ind',
        'contextualSpacing', 'mirrorMargins', 'textboxTightWrap', 'outlineLvl',
        'divId', 'cnvTxPreg', 'bidi', 'jc', 'textDirection', 'textAlignment',
        'textboxRect', 'rPr', 'sectPr', 'pPrChange'
    ]
    children = list(pPr)
    for child in children:
        pPr.remove(child)
    tag_to_index = {qn(f'w:{name}'): idx for idx, name in enumerate(order)}
    def get_key(child):
        return tag_to_index.get(child.tag, 999)
    children.sort(key=get_key)
    for child in children:
        pPr.append(child)

def sort_rPr(rPr):
    """Sort elements in run properties rPr strictly matching the Word OOXML schema."""
    order = [
        'rStyle', 'rFonts', 'b', 'bCs', 'i', 'iCs', 'caps', 'smallCaps', 'strike',
        'dstrike', 'outline', 'shadow', 'emboss', 'imprint', 'noProof', 'snapToGrid',
        'vanish', 'webHidden', 'color', 'spacing', 'w', 'position', 'sz', 'szCs',
        'highlight', 'u', 'effect', 'bdr', 'shd', 'fitText', 'vertAlign', 'rtl',
        'cs', 'em', 'specVanish', 'oMath', 'rPrChange'
    ]
    children = list(rPr)
    for child in children:
        rPr.remove(child)
    tag_to_index = {qn(f'w:{name}'): idx for idx, name in enumerate(order)}
    def get_key(child):
        return tag_to_index.get(child.tag, 999)
    children.sort(key=get_key)
    for child in children:
        rPr.append(child)

def format_paragraph(p, before=0, after=0, rtl=None, align=None, left_indent=None, right_indent=None, first_line_indent=None):
    """Set spacing, indents, alignment, and RTL bidirectional properties on a paragraph in a schema-compliant order."""
    pPr = p._p.get_or_add_pPr()
    if rtl is None:
        rtl = (pPr.find(qn('w:bidi')) is not None) or (p.alignment == WD_ALIGN_PARAGRAPH.RIGHT)
        
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.0
    
    if left_indent is not None:
        p.paragraph_format.left_indent = left_indent
    if right_indent is not None:
        p.paragraph_format.right_indent = right_indent
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
        
    if align is not None:
        p.alignment = align
    elif rtl:
        p.alignment = None
        
    if rtl:
        b = pPr.find(qn('w:bidi'))
        if b is None:
            b = OxmlElement('w:bidi')
            b.set(qn('w:val'), '1')
            pPr.append(b)
        if align is None:
            jc = pPr.find(qn('w:jc'))
            if jc is not None:
                pPr.remove(jc)
    else:
        b = pPr.find(qn('w:bidi'))
        if b is not None:
            pPr.remove(b)
            
    # Always clean up textDirection to prevent horizontal layout issues
    textDir = pPr.find(qn('w:textDirection'))
    if textDir is not None:
        pPr.remove(textDir)
        
    sort_pPr(pPr)

def para_space(para, before=0, after=0):
    format_paragraph(para, before=before, after=after)

def set_rtl_para(para):
    format_paragraph(para, rtl=True)

def set_rtl_run(run):
    rPr = run._r.get_or_add_rPr()
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
    sort_rPr(rPr)

def set_run_lang(run, lang_bidi="ar-SA"):
    rPr = run._r.get_or_add_rPr()
    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = OxmlElement('w:lang')
        rPr.append(lang)
    lang.set(qn('w:bidi'), lang_bidi)
    sort_rPr(rPr)

def make_run(para, text, bold=False, size=11, color=None, rtl=False, font_name='Arial'):
    # Coerce text to string and handle None/False cases safely
    if text is None or text is False:
        text = ""
    else:
        text = str(text)

    # Adapt default size 11 to fit the PDF's proportions (regular 10.0, bold 10.5)
    if size == 11:
        size = 10.5 if bold else 10.0
        
    # Auto-detect English/numbers inside RTL runs to override direction to LTR is disabled to maintain correct bidirectional order.
    pass
            
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font_name
    run.font.size = Pt(size)
    
    # Set complex script font size to match standard font size (Word uses half-points)
    rPr = run._r.get_or_add_rPr()
    szCs = rPr.find(qn('w:szCs'))
    if szCs is None:
        szCs = OxmlElement('w:szCs')
        rPr.append(szCs)
    szCs.set(qn('w:val'), str(int(size * 2)))
    
    # Set complex script bold to match standard bold
    if bold:
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            bCs = OxmlElement('w:bCs')
            rPr.append(bCs)
    
    if color:
        run.font.color.rgb = color
    if rtl:
        set_rtl_run(run)
        set_run_lang(run, "ar-SA")
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:cs'), 'Arial')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        sort_rPr(rPr)
    else:
        rPr = run._r.get_or_add_rPr()
        rtl_el = rPr.find(qn('w:rtl'))
        if rtl_el is not None:
            rPr.remove(rtl_el)
        sort_rPr(rPr)
    return run

def set_cell_shading(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)
    tcBorders = OxmlElement('w:tcBorders')
    for side, attrs in kwargs.items():
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), attrs.get('val','single'))
        el.set(qn('w:sz'), str(attrs.get('sz',6)))
        el.set(qn('w:space'),'0')
        el.set(qn('w:color'), attrs.get('color', NAVY_HEX))
        tcBorders.append(el)
    tcPr.append(tcBorders)

def clear_cell_borders(cell):
    set_cell_border(cell,
        top={'val':'none','sz':0,'color':'auto'},
        bottom={'val':'none','sz':0,'color':'auto'},
        left={'val':'none','sz':0,'color':'auto'},
        right={'val':'none','sz':0,'color':'auto'})

def kill_table_borders(table):
    tbl = table._tbl; tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'none'); el.set(qn('w:sz'),'0')
        el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto')
        tblBorders.append(el)
    tblPr.append(tblBorders)

def set_cell_valign(cell, align='top'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align); tcPr.append(vAlign)

def add_fld(para, fld_type):
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.text = f' {fld_type} '
    instrText.set(qn('xml:space'), 'preserve')
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run = para.add_run()
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def page_break_before(para):
    pPr = para._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore'); pPr.append(pb)

def set_cell_w(cell, w_cm):
    cell.width = w_cm
    tcPr = cell._tc.get_or_add_tcPr()
    existing_tcW = tcPr.find(qn('w:tcW'))
    if existing_tcW is not None:
        tcPr.remove(existing_tcW)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(w_cm.twips)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)

def set_col_widths(table, widths):
    tblPr = table._tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    total_width = sum(w.twips for w in widths)
    existing_tblW = tblPr.find(qn('w:tblW'))
    if existing_tblW is not None:
        tblPr.remove(existing_tblW)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(total_width)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_w(row.cells[idx], width)


NAVY_BORDER      = {'val':'single','sz':6,'color':NAVY_HEX}
NO_BORDER        = {'val':'none','sz':0,'color':'auto'}
NAVY_GRID_BORDER = {'val':'single','sz':4,'color':NAVY_HEX}

# ── Build Document ────────────────────────────────────────────────────────
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11.5)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after  = Pt(0)
style.paragraph_format.line_spacing = 1.0

sec = doc.sections[0]
sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
sec.left_margin = Cm(0.8); sec.right_margin = Cm(0.8)
sec.header_distance = Cm(0.8); sec.footer_distance = Cm(0.8)
PAGE_W = 19.4  # usable cm

# Section 1 has only one page (cover page), so we show the default header/footer
sec.different_first_page_header_footer = False

# ═══════════════════════════════════════════════════════════════
# COVER PAGE (Page 1)
# ═══════════════════════════════════════════════════════════════

# Header row: EN | Logo | AR (Default header for Section 1 since it's 1 page)
hdr1 = sec.header
t_hdr = hdr1.add_table(rows=1, cols=3, width=Cm(PAGE_W))
p0 = hdr1.paragraphs[0]
p0._element.getparent().remove(p0._element)
t_hdr.allow_autofit = False
kill_table_borders(t_hdr)
for i, c in enumerate(t_hdr.row_cells(0)):
    clear_cell_borders(c)
    set_cell_border(c, bottom=NAVY_BORDER)
    set_cell_valign(c, 'top')
set_col_widths(t_hdr, [Cm(8.2), Cm(3.0), Cm(8.2)])

# Set padding to zero on extreme sides
set_cell_margins(t_hdr.cell(0, 0), top=0, bottom=0, left=0)
set_cell_margins(t_hdr.cell(0, 1), top=0, bottom=0, left=0, right=0)
set_cell_margins(t_hdr.cell(0, 2), top=0, bottom=0, right=0)

p1 = t_hdr.cell(0,0).paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p1, before=0, after=0)
make_run(p1, 'Hussein and Al Hassan G. Shaker Bros.', bold=True, size=10.5, color=NAVY)
p2 = t_hdr.cell(0,0).add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p2, before=0, after=0)
make_run(p2, 'For Modern Trading Co. Ltd.', bold=True, size=10.5, color=NAVY)

p = t_hdr.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_paragraph(p, before=0, after=0)
if os.path.exists(LOGO_PATH):
    p.add_run().add_picture(LOGO_PATH, width=Cm(2.6))

p1_ar = t_hdr.cell(0,2).paragraphs[0]; set_rtl_para(p1_ar)
format_paragraph(p1_ar, before=0, after=0, rtl=True)
make_run(p1_ar, 'شركة الأخوان حسين و الحسن غازي شاكر', bold=True, size=10.5, color=NAVY, rtl=True)
p2_ar = t_hdr.cell(0,2).add_paragraph(); set_rtl_para(p2_ar)
format_paragraph(p2_ar, before=0, after=0, rtl=True)
make_run(p2_ar, 'للتجارة الحديثة المحدودة', bold=True, size=10.5, color=NAVY, rtl=True)

_sp1 = hdr1.add_paragraph(); para_space(_sp1, 0, 0)

# Date & Reference Table
t = doc.add_table(rows=2, cols=3)
t.allow_autofit = False
kill_table_borders(t)
for r in t.rows:
    for c in r.cells: clear_cell_borders(c)
set_col_widths(t, [Cm(5.0), Cm(9.4), Cm(5.0)])

for r_idx in range(2):
    set_cell_margins(t.cell(r_idx, 0), top=0, bottom=0, left=0)
    set_cell_margins(t.cell(r_idx, 1), top=0, bottom=0, left=0, right=0)
    set_cell_margins(t.cell(r_idx, 2), top=0, bottom=0, right=0)

# Row 1: Date
p = t.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p, before=0, after=0)
make_run(p, 'Date', bold=True, size=12, color=NAVY)

p = t.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_paragraph(p, before=0, after=0)
make_run(p, today_str, bold=True, size=12, color=NAVY)

p = t.cell(0,2).paragraphs[0]; set_rtl_para(p)
format_paragraph(p, before=0, after=0, rtl=True)
make_run(p, 'التاريخ', bold=True, size=12, color=NAVY, rtl=True)

# Row 2: Reference
p = t.cell(1,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p, before=0, after=0)
make_run(p, 'Reference No.', bold=True, size=12, color=NAVY)

p = t.cell(1,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_paragraph(p, before=0, after=0)
make_run(p, ref_no, bold=True, size=12, color=NAVY)

p = t.cell(1,2).paragraphs[0]; set_rtl_para(p)
format_paragraph(p, before=0, after=0, rtl=True)
make_run(p, 'الرقم المرجعي', bold=True, size=12, color=NAVY, rtl=True)

_sp = doc.add_paragraph(); para_space(_sp, 4, 4)

t = doc.add_table(rows=1, cols=2)
t.allow_autofit = False
kill_table_borders(t)
for c in t.row_cells(0): clear_cell_borders(c)
set_col_widths(t, [Cm(9.7), Cm(9.7)])

set_cell_margins(t.cell(0, 0), top=0, bottom=0, left=0)
set_cell_margins(t.cell(0, 1), top=0, bottom=0, right=0)

p = t.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p, before=0, after=0)
make_run(p, 'Planned Service Agreement', bold=True, size=16, color=NAVY)

p = t.cell(0,1).paragraphs[0]; set_rtl_para(p)
format_paragraph(p, before=0, after=0, rtl=True)
make_run(p, 'اتفاقية خدمة مجدولة', bold=True, size=16, color=NAVY, rtl=True)

# Cover image
if os.path.exists(COVER_IMG_PATH):
    pi = doc.add_paragraph(); pi.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pi.add_run().add_picture(COVER_IMG_PATH, width=Cm(PAGE_W))

# Push contact info to bottom of page
spacer = doc.add_paragraph()
para_space(spacer, before=180, after=0)

# Contact info with icons matching PDF layout
ICON_SIZE = Cm(0.8)  # Height matching PDF
PHONE_PATH   = os.path.join(IMG_DIR, 'phone.png')
WORLD_PATH   = os.path.join(IMG_DIR, 'world.png')
FB_PATH      = os.path.join(IMG_DIR, 'facebook.png')

contact_data = [
    (PHONE_PATH,  '8002440247'),
    (WORLD_PATH,  'HH-SHAKER.COM.SA'),
    (FB_PATH,     'HHSHAKERCO'),
]
for icon_path, txt in contact_data:
    p = doc.add_paragraph()
    para_space(p, before=4, after=12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Add icon image inline
    if os.path.exists(icon_path):
        run_img = p.add_run()
        run_img.add_picture(icon_path, height=ICON_SIZE)
    # Add text next to icon
    make_run(p, '  ' + txt, bold=True, size=12.5, color=NAVY)

def build_text_footer(f_sec, include_page_num=True):
    f_sec.is_linked_to_previous = False
    tf = f_sec.add_table(rows=1, cols=3, width=Cm(PAGE_W))
    p0 = f_sec.paragraphs[0]
    p0._element.getparent().remove(p0._element)
    tf.allow_autofit = False
    kill_table_borders(tf)
    set_col_widths(tf, [Cm(8.2), Cm(3.0), Cm(8.2)])
    for i, c in enumerate(tf.row_cells(0)):
        clear_cell_borders(c)
        set_cell_border(c, top=NAVY_BORDER)
        if i == 0:
            set_cell_margins(c, top=0, bottom=0, left=0)
        elif i == 2:
            set_cell_margins(c, top=0, bottom=0, right=0)
        else:
            set_cell_margins(c, top=0, bottom=0, left=0, right=0)
    p1 = tf.cell(0,0).paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(p1, before=0, after=0)
    make_run(p1, 'Hussein and Al Hassan G. Shaker Bros.', bold=True, size=10.5, color=NAVY)
    p2 = tf.cell(0,0).add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    format_paragraph(p2, before=0, after=0)
    make_run(p2, 'For Modern Trading Co. Ltd.', bold=True, size=10.5, color=NAVY)
    
    p = tf.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(p, before=0, after=0)
    if include_page_num:
        make_run(p, 'Page ', bold=True, size=10.5, color=NAVY); add_fld(p, 'PAGE')
        make_run(p, ' of ', bold=True, size=10.5, color=NAVY); add_fld(p, 'SECTIONPAGES')
        
    p1_ar = tf.cell(0,2).paragraphs[0]; set_rtl_para(p1_ar)
    format_paragraph(p1_ar, before=0, after=0, rtl=True)
    make_run(p1_ar, 'شركة الأخوان حسين و الحسن غازي شاكر', bold=True, size=10.5, color=NAVY, rtl=True)
    p2_ar = tf.cell(0,2).add_paragraph(); set_rtl_para(p2_ar)
    format_paragraph(p2_ar, before=0, after=0, rtl=True)
    make_run(p2_ar, 'للتجارة الحديثة المحدودة', bold=True, size=10.5, color=NAVY, rtl=True)
    _spf = f_sec.add_paragraph(); para_space(_spf, 0, 0)

# Footer for page 1 (cover page footer - default footer for Section 1 since titlePg is not used)
build_text_footer(sec.footer, include_page_num=False)

# ═══════════════════════════════════════════════════════════════
# Page break to content pages
# ═══════════════════════════════════════════════════════════════
def add_explicit_page_break():
    bp = doc.add_paragraph()
    pPr = bp._p.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    pPr.append(spacing)
    run_brk = bp.add_run()
    brk = OxmlElement('w:br'); brk.set(qn('w:type'), 'page')
    run_brk._r.append(brk)

# Section 2: Content Pages
sec2 = doc.add_section(docx.enum.section.WD_SECTION_START.NEW_PAGE)
sec2.top_margin = Cm(1.5); sec2.bottom_margin = Cm(1.5)
sec2.left_margin = Cm(0.8); sec2.right_margin = Cm(0.8)
sec2.header_distance = Cm(0.8); sec2.footer_distance = Cm(0.8)

# Restart page numbering at 1 in Section 2
sec2_pr = sec2._sectPr
pgNumType = OxmlElement('w:pgNumType')
pgNumType.set(qn('w:start'), '1')
sec2_pr.append(pgNumType)

# Unlink headers/footers of Section 2 from Section 1
sec2.header.is_linked_to_previous = False
sec2.footer.is_linked_to_previous = False

# Header for pages 2+
th = sec2.header.add_table(rows=1, cols=3, width=Cm(PAGE_W))
p0 = sec2.header.paragraphs[0]
p0._element.getparent().remove(p0._element)
th.allow_autofit = False
kill_table_borders(th)
set_col_widths(th, [Cm(8.2), Cm(3.0), Cm(8.2)])
for i, c in enumerate(th.row_cells(0)):
    clear_cell_borders(c)
    set_cell_border(c, bottom=NAVY_BORDER)
    set_cell_valign(c, 'top')
    if i == 0:
        set_cell_margins(c, top=0, bottom=0, left=0)
    elif i == 2:
        set_cell_margins(c, top=0, bottom=0, right=0)
    else:
        set_cell_margins(c, top=0, bottom=0, left=0, right=0)
p = th.cell(0,0).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
format_paragraph(p, before=0, after=0)
make_run(p, 'Planned Service Agreement', bold=True, size=10.5, color=NAVY)
p = th.cell(0,1).paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
format_paragraph(p, before=0, after=0)
if os.path.exists(LOGO_PATH):
    p.add_run().add_picture(LOGO_PATH, width=Cm(2.6))
p = th.cell(0,2).paragraphs[0]; set_rtl_para(p)
format_paragraph(p, before=0, after=0, rtl=True)
make_run(p, 'اتفاقية الخدمة المخططة', bold=True, size=10.5, color=NAVY, rtl=True)
_sph = sec2.header.add_paragraph(); para_space(_sph, 0, 0)

# Footer for pages 2+
build_text_footer(sec2.footer, include_page_num=True)

# ═══════════════════════════════════════════════════════════════
# BILINGUAL CONTENT TABLE
# ═══════════════════════════════════════════════════════════════
current_table = None

def start_new_bilingual_table():
    global current_table
    current_table = doc.add_table(rows=0, cols=2)
    current_table.allow_autofit = False
    kill_table_borders(current_table)
    
    tblPr = current_table._tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(Cm(PAGE_W).twips)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

def close_current_bilingual_table():
    if current_table and len(current_table.rows) > 0:
        last_row = current_table.rows[-1]
        set_cell_border(last_row.cells[0], left=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)
        set_cell_border(last_row.cells[1], right=NAVY_BORDER, bottom=NAVY_BORDER)

def insert_page_break():
    close_current_bilingual_table()
    add_explicit_page_break()
    start_new_bilingual_table()

start_new_bilingual_table()

def add_row(en_fn, ar_fn):
    row = current_table.add_row()
    en_c = row.cells[0]; ar_c = row.cells[1]
    w_val = Cm(PAGE_W / 2)
    set_cell_w(en_c, w_val)
    set_cell_w(ar_c, w_val)
    for p in en_c.paragraphs: p.clear()
    for p in ar_c.paragraphs: p.clear()
    clear_cell_borders(en_c); clear_cell_borders(ar_c)
    
    if len(current_table.rows) == 1:
        # First row of this table needs top border too
        set_cell_border(en_c, left=NAVY_BORDER, right=NAVY_BORDER, top=NAVY_BORDER)
        set_cell_border(ar_c, right=NAVY_BORDER, top=NAVY_BORDER)
    else:
        set_cell_border(en_c, left=NAVY_BORDER, right=NAVY_BORDER)
        set_cell_border(ar_c, right=NAVY_BORDER)
        
    set_cell_valign(en_c, 'top'); set_cell_valign(ar_c, 'top')
    en_fn(en_c); ar_fn(ar_c)

def title_heading(en, ar):
    def fn_en(c):
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_space(p, 1, 0)
        make_run(p, en, bold=True, size=16, color=NAVY)
    def fn_ar(c):
        p = c.paragraphs[0]; set_rtl_para(p)
        para_space(p, 1, 0)
        make_run(p, ar, bold=True, size=16, color=NAVY, rtl=True)
    add_row(fn_en, fn_ar)

def heading(en, ar, indent=0):
    def fn_en(c):
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        if indent: p.paragraph_format.left_indent = Cm(indent)
        para_space(p, 1, 0)
        make_run(p, en, bold=True, size=14, color=NAVY)
    def fn_ar(c):
        p = c.paragraphs[0]; set_rtl_para(p)
        if indent: p.paragraph_format.right_indent = Cm(indent)
        para_space(p, 1, 0)
        make_run(p, ar, bold=True, size=14, color=NAVY, rtl=True)
    add_row(fn_en, fn_ar)

# Insert the missing top headings inside the table
title_heading('Planned Service Agreement', 'اتفاقية الخدمة المخططة')
heading('Parties', 'الأطراف')

def text_row(en, ar):
    def fn_en(c):
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para_space(p, 0, 0); make_run(p, en, size=11)
    def fn_ar(c):
        p = c.paragraphs[0]; set_rtl_para(p)
        para_space(p, 0, 0); make_run(p, ar, size=11, rtl=True)
    add_row(fn_en, fn_ar)

def bullets(en_items, ar_items):
    def fn_en(c):
        first = True
        for item in en_items:
            p = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.left_indent = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(-0.3)
            para_space(p, 0, 0)
            make_run(p, u'\u2022  ', size=11)
            if isinstance(item, list):
                for txt, bld in item:
                    make_run(p, txt, bold=bld, size=11)
            else:
                make_run(p, item, size=11)
    def fn_ar(c):
        first = True
        for item in ar_items:
            p = c.paragraphs[0] if first else c.add_paragraph()
            first = False
            set_rtl_para(p)
            p.paragraph_format.right_indent  = Cm(1.0)
            p.paragraph_format.first_line_indent = Cm(0.3)   # positive = bullet hangs into right margin
            para_space(p, 0, 0)
            make_run(p, u'\u2022  ', size=11, rtl=True)
            if isinstance(item, list):
                for txt, bld in item:
                    make_run(p, txt, bold=bld, size=11, rtl=True)
            else:
                make_run(p, item, size=11, rtl=True)
    add_row(fn_en, fn_ar)

def numbered(num, en_lbl, en_txt, ar_lbl, ar_txt, sub_en=None, sub_ar=None):
    def fn_en(c):
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.left_indent = Cm(0.3)
        p.paragraph_format.first_line_indent = Cm(-0.3)
        para_space(p, 0, 0)
        make_run(p, f'{num}. ', bold=True, size=11, color=NAVY)
        make_run(p, en_lbl, bold=True, size=11, color=NAVY)
        make_run(p, en_txt, size=11)
        if sub_en:
            for s in sub_en:
                sp = c.add_paragraph()
                sp.paragraph_format.left_indent = Cm(1.0)
                sp.paragraph_format.first_line_indent = Cm(-0.3)
                para_space(sp, 0, 0)
                make_run(sp, u'\u2022  ' + s, size=11)
    def fn_ar(c):
        p = c.paragraphs[0]; set_rtl_para(p)
        p.paragraph_format.right_indent = Cm(0.3)
        p.paragraph_format.first_line_indent = Cm(0.3)   # positive = number hangs into right margin
        para_space(p, 0, 0)
        make_run(p, f'{num}. ', bold=True, size=11, color=NAVY, rtl=True)
        make_run(p, ar_lbl, bold=True, size=11, color=NAVY, rtl=True)
        make_run(p, ar_txt, size=11, rtl=True)
        if sub_ar:
            for s in sub_ar:
                sp = c.add_paragraph()
                set_rtl_para(sp)
                sp.paragraph_format.right_indent = Cm(1.0)
                sp.paragraph_format.first_line_indent = Cm(0.3)   # positive = bullet hangs into right margin
                para_space(sp, 0, 0)
                make_run(sp, u'\u2022  ' + s, size=11, rtl=True)
    add_row(fn_en, fn_ar)

# ── First Party ──
def fn1e(c):
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; para_space(p, 1, 1)
    make_run(p, 'First Party: ', bold=True, size=11, color=NAVY)
    make_run(p, 'Hussein and Al Hassan G. Shaker Bros. for Modern Trading Co. Ltd.', bold=True, size=11)
    make_run(p, ' represented in this contract by ', size=11)
    make_run(p, 'Mr.Nouraldeen Riyad Nofal', bold=True, size=11)
    make_run(p, ' ID / Resident No. ', size=11)
    make_run(p, '2396646263', bold=True, size=11)
    make_run(p, ' in his capacity as ', size=11)
    make_run(p, 'National Service Manager', bold=True, size=11)
    make_run(p, ' Mobile No. ', size=11)
    make_run(p, '+966550510416', bold=True, size=11)
    make_run(p, ' and Email address ', size=11)
    make_run(p, 'nouraldeen@hh-shaker.com.sa', bold=True, size=11)
    make_run(p, ' (Hereinafter referred to as the “First Party” and in short HH Shaker ServicePro)', size=11)
def fn1a(c):
    p = c.paragraphs[0]; set_rtl_para(p); para_space(p, 1, 1)
    make_run(p, 'الطرف الأول: ', bold=True, size=11, color=NAVY, rtl=True)
    make_run(p, 'شركة الإخوان حسين والحسن غازي شاكر للتجارة الحديثة المحدودة  ', bold=True, size=11, rtl=True)
    make_run(p, 'ويمثلها في هذا العقد  ', size=11, rtl=True)
    make_run(p, 'Mr.Nouraldeen Riyad Nofal', bold=True, size=11, rtl=True)
    make_run(p, ' رقم الهوية / الإقامة  ', size=11, rtl=True)
    make_run(p, '2396646263', bold=True, size=11, rtl=True)
    make_run(p, ' بصفته  ', size=11, rtl=True)
    make_run(p, 'National Service Manager', bold=True, size=11, rtl=True)
    make_run(p, ' رقم الجوال  ', size=11, rtl=True)
    make_run(p, '+966550510416', bold=True, size=11, rtl=True)
    make_run(p, ' والبريد الإلكتروني  ', size=11, rtl=True)
    make_run(p, 'nouraldeen@hh-shaker.com.sa', bold=True, size=11, rtl=True)
    make_run(p, ' (ويشار إليه فيما بعد بـ "الطرف الأول" أو باختصار HH Shaker ServicePro)', size=11, rtl=True)
add_row(fn1e, fn1a)

# ── Second Party ──
def fn2e(c):
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT; para_space(p, 1, 1)
    make_run(p, 'Second Party: ', bold=True, size=11, color=NAVY)
    make_run(p, partner_name, bold=True, size=11)
    make_run(p, ' represented in this contract by ', size=11)
    make_run(p, customer_name, bold=True, size=11)
    make_run(p, ' ID / Resident No. ', size=11)
    make_run(p, id_party, bold=True, size=11)
    make_run(p, ' in his capacity as ', size=11)
    make_run(p, job_position, bold=True, size=11)
    make_run(p, ' Mobile No. ', size=11)
    make_run(p, mobile_no, bold=True, size=11)
    make_run(p, ' Email: ', size=11)
    make_run(p, email_addr, bold=True, size=11)
    make_run(p, ' (Hereinafter referred to as the “Second Party”)', size=11)
def fn2a(c):
    p = c.paragraphs[0]; set_rtl_para(p); para_space(p, 1, 1)
    make_run(p, 'الطرف الثاني: ', bold=True, size=11, color=NAVY, rtl=True)
    make_run(p, partner_name, bold=True, size=11, rtl=True)
    make_run(p, ' ويمثله في هذا العقد ', size=11, rtl=True)
    make_run(p, customer_name, bold=True, size=11, rtl=True)
    make_run(p, ' رقم الهوية / الإقامة ', size=11, rtl=True)
    make_run(p, id_party, bold=True, size=11, rtl=True)
    make_run(p, ' بصفته ', size=11, rtl=True)
    make_run(p, job_position, bold=True, size=11, rtl=True)
    make_run(p, ' رقم الجوال ', size=11, rtl=True)
    make_run(p, mobile_no, bold=True, size=11, rtl=True)
    make_run(p, ' والبريد الإلكتروني ', size=11, rtl=True)
    make_run(p, email_addr, bold=True, size=11, rtl=True)
    make_run(p, ' (ويشار إليه فيما بعد بـ "الطرف الثاني")', size=11, rtl=True)
add_row(fn2e, fn2a)
text_row('Collectively referred to as the “Parties”.', 'يشار إليها مجتمعة باسم "الأطراف".')

# ── Section 1: Introduction ──
heading('1. Introduction', '1. مقدمة')
text_row(
    'Whereas the first party possesses the technical expertise, specialized team and tools necessary for the maintenance of air conditioning systems. Whereas the second party possesses facilities that require periodic maintenance and repair work and needs to benefit from the first party\'s services in this field, and the first party has agreed to this, the will of the two parties, who are in full legal and statutory capacity, has converged to conclude this contract and its following terms.',
    'وحيث أن الطرف الأول يمتلك الخبرة الفنية والفريق المتخصص والأدوات اللازمة لصيانة أنظمة التكييف. وحيث يمتلك الطرف الثاني مرافق تتطلب صيانة دورية واعمال الصلاح ويرغب في الإنتفاع من خدمات الطرف الأول في هذا المجال ووافق الطرف الأول على ذلك فقد إلتقت إرادة الطرفين وهم بكامل أهليتهم الشرعية والنظامية على إبرام هذا العقد وبنوده التالية.')

# ── Section 2: Contract Type ──
heading('2. Contract Type and Service Coverage', '2. نوع العقد وتغطية الخدمة')
heading('2.1 Periodic Preventive Maintenance', '2.1 الصيانة الوقائية الدورية', indent=0.8)

for line in seen_visits.values():
    nv = line.get('no_of_visits_per_year', 0) or 0
    pn = line.get('product_name', '')
    pa = line.get('product_arabic_name', '')
    wn = line.get('visits_words_en', str(nv))
    wa = line.get('visits_words_ar', str(nv))
    bullets(
        [[(wn, True), (' (', False), (str(nv), True), (') scheduled visits per year. (', False), (pn, True), (')', False)]],
        [[(wa, True), (' (', False), (str(nv), True), (') زيارات مجدولة سنوياً. (', False), (pa, True), (')', False)]]
    )

bullets(['All works to be completed within 12 months from the contract start date.'],
        ['سيتم الانتهاء من جميع الأعمال خلال 12 شهراً من تاريخ بدء العقد.'])
bullets(['Detailed scope of preventive maintenance is listed in Annex A – Schedule of Preventive Tasks.'],
        ['ويرد النطاق التفصيلي للصيانة الوقائية في الملحق أ - جدول المهام الوقائية.'])
bullets(['Details of the site, equipment and its quantity are listed in Annex B'],
        ['ترد تفاصيل الموقع والمعدات والكميات في المرفق ب'])

heading('2.2 Emergency Callout Service', '2.2 خدمة نداءات الطوارئ', indent=0.8)
bullets(['Response within 24 working hours (maximum).'],
        ['الاستجابة خلال 24 ساعة عمل كحد أقصى.'])
bullets([[('Includes ', False), (f'{total_emergency} emergency visits', True), (' annually.', False)]],
        [[('يشمل ', False), (f'{total_emergency} زيارة طارئة', True), (' سنوياً.', False)]])
insert_page_break()
bullets(['Non Comprehensive contracts covers repairs that can be completed within 3 working hours without any additional spare parts.'],
        ['العقود غير الشاملة تشمل الإصلاحات التي يمكن إنجازها خلال 3 ساعات عمل دون الحاجة إلى قطع غيار إضافية.'])
bullets(['Semi Comprehensive contracts covers repairs that can be completed within 3 working hours with additional spare parts except compressor and coils.'],
        ['العقود شبه الشاملة تشمل الإصلاحات التي يمكن إنجازها خلال 3 ساعات عمل مع توفير قطع غيار إضافية باستثناء الكمبروسر وملفاته.'])
bullets(['Full Comprehensive contract covers repairs that can be completed within 3 working hours with additional spare parts.'],
        ['العقود الكاملة تشمل الإصلاحات التي يمكن إنجازها خلال 3 ساعات عمل مع توفير قطع غيار إضافية.'])
bullets([[('Additional visits charged at ', False), (f'{add_paid_price:,.2f}', True), (' per visit (including VAT).', False)]],
        [[('الزيارات الإضافية بقيمة ', False), (f'{add_paid_price:,.2f}', True), (' ريال سعودي لكل زيارة (بما في ذلك ضريبة القيمة المضافة).', False)]])

heading('2.3 Major Corrective Maintenance', '2.3 الصيانة التصحيحية الرئيسية', indent=0.8)
bullets(['Covers repair works requiring spare parts (e.g., compressor, condenser coil, evaporator coil, blower assembly, leakage repairs).'],
        ['تشمل أعمال الإصلاح التي تتطلب قطع غيار (على سبيل المثال: الكمبروسر، ملف المكثف، ملف المبخر، مجموعة المنفاخ، إصلاح التسرب).'])
bullets(['For Non Comprehensive contracts, a separate quotation including parts and labor will be issued for each corrective repair.'],
        ['بالنسبة للعقود غير الشاملة، سيتم إصدار عرض سعر منفصل يشمل قطع الغيار وأجور اليد العاملة لكل عملية إصلاح.'])
bullets(['For Semi Comprehensive contracts, a separate quotation including parts and labor will be issued for each corrective repair of compressor and coils.'],
        ['بالنسبة للعقود شبه الشاملة، سيتم إصدار عرض سعر منفصل يشمل قطع الغيار وأجور اليد العاملة لكل عملية إصلاح الكمبروسر والملفات.'])
bullets(['For Full comprehensive contract, a separate quotation including parts and labor will be issued for replacement of coils. (Only Package & VRF Outdoor coils are covered under comprehensive contract)'],
        ['بالنسبة للعقود الكاملة، سيتم إصدار عرض سعر منفصل يشمل قطع الغيار وأجور اليد العاملة لاستبدال الملفات. (فقط وحدات الباكيج ووحدات VRF الخارجية مشمولة في العقد الشامل)'])
bullets(['A technical report will be provided after each visit'], ['سيتم تقديم تقرير فني بعد كل زيارة.'])

# ── Section 3: Obligations ──
heading('3. Parties\' Obligations', '3. التزامات الأطراف')
heading('3.1 First Party\'s Obligations', '3.1 التزامات الطرف الأول', indent=0.8)
bullets(['Supervise technicians and ensure service quality.'], ['الإشراف على الفنيين والتأكد من جودة الخدمة.'])
bullets(['Perform maintenance in accordance with required standards'], ['إجراء الصيانة وفقاً للمعايير المطلوبة.'])
bullets(['Adhere to agreed maintenance dates.'], ['الالتزام بتواريخ الصيانة المتفق عليها.'])
bullets(['Avoid damage to facilities or equipment of the Second Party.'], ['تجنب إلحاق الضرر بمرافق أو معدات الطرف الثاني.'])

heading('3.2 Second Party\'s Obligations', '3.2 التزامات الطرف الثاني', indent=0.8)
bullets(['Provide access and necessary resources (power, water, permits, safe access).'], ['توفير الوصول والموارد اللازمة (الكهرباء والمياه والتصاريح والوصول الآمن).'])
bullets(['Pay dues as per the agreed terms.'], ['دفع المستحقات حسب الشروط المتفق عليها.'])
insert_page_break()
bullets(['Pay the amount of covering additional visits or services outside the scope of the contract.'], ['سداد قيمة تغطية الزيارات أو الخدمات الإضافية خارج نطاق العقد.'])
bullets(['Issue purchase orders (POs) for corrective works as needed.'], ['إصدار أوامر الشراء (POs) للأعمال التصحيحية عند الحاجة.'])
bullets(['Ensure units are in safe and operable condition before contract commencement.'], ['التأكد من أن جميع وحدات التكييف في حالة آمنة وقابلة للتشغيل قبل بدء العقد.'])

# ── Section 4: Duration ──
heading('4. Duration of Agreement', '4. مدة الاتفاقية')
def fn4e(c):
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(1.5); p.paragraph_format.first_line_indent = Cm(-0.4)
    para_space(p, 0, 0)
    make_run(p, u'\u2022  This Agreement shall commence on ', size=11)
    make_run(p, date_start, bold=True, size=11)
    make_run(p, ' and end on ', size=11)
    make_run(p, date_end, bold=True, size=11)
    make_run(p, '.', size=11)
    for txt in ['Automatically renewable for one (1) year unless either Party gives written notice at least two (2) months before expiry.',
                'Agreement value may be increased by 5\u201310% annually upon mutual approval.']:
        p2 = c.add_paragraph(); p2.paragraph_format.left_indent = Cm(1.5)
        p2.paragraph_format.first_line_indent = Cm(-0.4); para_space(p2, 0, 0)
        make_run(p2, u'\u2022  ' + txt, size=11)
def fn4a(c):
    p = c.paragraphs[0]; set_rtl_para(p)
    p.paragraph_format.right_indent = Cm(1.5); p.paragraph_format.first_line_indent = Cm(0.4)
    para_space(p, 0, 0)
    make_run(p, u'\u2022  تبدأ هذه الاتفاقية في ', size=11, rtl=True)
    make_run(p, date_start, bold=True, size=11, rtl=True)
    make_run(p, ' وتنتهي في ', size=11, rtl=True)
    make_run(p, date_end, bold=True, size=11, rtl=True)
    make_run(p, '.', size=11, rtl=True)
    for txt in ['قابلة للتجديد تلقائياً لمدة عام واحد (1) ما لم يقدم أي من الطرفين إشعاراً كتابياً قبل شهرين (2) على الأقل من انتهاء الصلاحية.',
                'يمكن زيادة قيمة الاتفاقية بنسبة 5\u201310٪ سنوياً بموافقة الطرفين.']:
        p2 = c.add_paragraph(); set_rtl_para(p2)
        p2.paragraph_format.right_indent = Cm(1.5); p2.paragraph_format.first_line_indent = Cm(0.4)
        para_space(p2, 0, 0); make_run(p2, u'\u2022  ' + txt, size=11, rtl=True)
add_row(fn4e, fn4a)

# ── Section 5: Payment ──
heading('5. Payment Terms', '5. شروط الدفع')
def fn5e(c):
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.first_line_indent = Cm(-0.3)
    para_space(p, 0, 0)
    make_run(p, u'\u2022  Annual contract value: ', size=11)
    make_run(p, f'SAR {amount_total:,.2f}', bold=True, size=11)
    make_run(p, ' (including VAT).', size=11)
    p2 = c.add_paragraph(); p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.first_line_indent = Cm(-0.3); para_space(p2, 0, 0)
    make_run(p2, u'\u2022  Payment shall be made ', size=11)
    make_run(p2, payment_term_text, bold=True, size=11)
    make_run(p2, '.', size=11)
    for s in payment_schedule_lines:
        ps = c.add_paragraph(); ps.paragraph_format.left_indent = Cm(1.0)
        ps.paragraph_format.first_line_indent = Cm(-0.3); para_space(ps, 0, 0)
        make_run(ps, f"\u2013  ", size=11)
        make_run(ps, f"{s.get('name','')}", bold=True, size=11)
        make_run(ps, f" \u2013 ", size=11)
        make_run(ps, f"{s.get('payment_date','')}", bold=True, size=11)
        make_run(ps, f" \u2013 ", size=11)
        make_run(ps, f"{s.get('amount',0):,.2f} SR", bold=True, size=11)
    p3 = c.add_paragraph(); p3.paragraph_format.left_indent = Cm(0.5)
    p3.paragraph_format.first_line_indent = Cm(-0.3); para_space(p3, 0, 0)
    make_run(p3, u'\u2022  Late payments may result in the suspension of services until the overdue dues are settled, and the first party has the right to terminate the contract if the second party does not commit to paying any of the financial dues.', size=11)
def fn5a(c):
    p = c.paragraphs[0]; set_rtl_para(p)
    p.paragraph_format.right_indent = Cm(0.5); p.paragraph_format.first_line_indent = Cm(0.3)
    para_space(p, 0, 0)
    make_run(p, u'\u2022  قيمة العقد السنوي: ', size=11, rtl=True)
    make_run(p, f'{amount_total:,.2f} ريال سعودي', bold=True, size=11, rtl=True)
    make_run(p, ' (شامل ضريبة القيمة المضافة).', size=11, rtl=True)
    p2 = c.add_paragraph(); set_rtl_para(p2)
    p2.paragraph_format.right_indent = Cm(0.5); p2.paragraph_format.first_line_indent = Cm(0.3)
    para_space(p2, 0, 0)
    make_run(p2, u'\u2022  يتم الدفع ', size=11, rtl=True)
    make_run(p2, payment_term_text_ar, bold=True, size=11, rtl=True)
    make_run(p2, '.', size=11, rtl=True)
    for s in payment_schedule_lines:
        ps = c.add_paragraph(); set_rtl_para(ps)
        ps.paragraph_format.right_indent = Cm(1.0); ps.paragraph_format.first_line_indent = Cm(0.3)
        para_space(ps, 0, 0)
        make_run(ps, f"\u2013  ", size=11, rtl=True)
        make_run(ps, f"{s.get('name_ara','')}", bold=True, size=11, rtl=True)
        make_run(ps, f" \u2013 ", size=11, rtl=True)
        make_run(ps, f"{s.get('payment_date','')}", bold=True, size=11, rtl=True)
        make_run(ps, f" \u2013 ", size=11, rtl=True)
        make_run(ps, f"{s.get('amount',0):,.2f} س.ر", bold=True, size=11, rtl=True)
    p3 = c.add_paragraph(); set_rtl_para(p3)
    p3.paragraph_format.right_indent = Cm(0.5); p3.paragraph_format.first_line_indent = Cm(0.3)
    para_space(p3, 0, 0)
    make_run(p3, u'\u2022  قد يؤدي التأخر في المدفوعات إلى تعليق الخدمات حتى يتم تسوية المستحقات المتأخرة ويحق للطرف الأول إنهاء العقد في حال لم يلتزم الطرف الثاني بسداد أي من المستحقات المالية.', size=11, rtl=True)
add_row(fn5e, fn5a)

# ── Section 6: General Terms ──
heading('6. General Terms & Conditions', '6. الشروط والأحكام العامة')
numbered('1', 'Service Hours: ', '8:30 AM \u2013 5:30 PM, Saturday to Thursday (excluding Fridays and public holidays). Work outside these hours will be charged separately and determined by the first party.',
         'ساعات الخدمة: ', '8:30 صباحاً \u2013 5:30 مساءً، من السبت إلى الخميس (باستثناء الجمعة والعطلات الرسمية). العمل خارج هذه الساعات يُحتسب بشكل منفصل ويحدده الطرف الأول.')
numbered('2', 'Safety & Site Access: ', 'The Second Party must provide safe access and comply with security/safety standards.',
         'السلامة والوصول إلى الموقع: ', 'يجب على الطرف الثاني توفير الوصول الآمن والامتثال لمعايير الأمن / السلامة.')
numbered('3', 'Repairs & Spare Parts: ', 'Only First Party technicians are authorized for covered units. If serviced by third party, First Party holds no liability.',
         'الإصلاحات وقطع الغيار: ', 'فنيو الطرف الأول فقط مصرح لهم بالوحدات المغطاة بالعقد. إذا تمت الخدمة من قبل الطرف ثالث، فإن الطرف الأول لا يتحمل أي مسؤولية.',
         sub_en=['As we are the exclusive agent for Midea Air conditioners, it is recommended that spare parts be purchased from the First Party.'],
         sub_ar=['بما أننا الوكيل الحصري لمكيفات Midea، يُوصى بشراء قطع الغيار من الطرف الأول.'])
insert_page_break()
numbered('4', 'Limitation of Liability: ', '',
         'حدود المسؤولية: ', '',
         sub_en=['First Party shall not be responsible for delays or damages caused by force majeure (e.g., natural disasters, strikes, fires, unavailability of parts).',
                 'Liability does not cover units not included in this contract.',
                 'If force majeure prevents execution for four (4) months or more, either Party may terminate the agreement by written notice.'],
         sub_ar=['لا يتحمل الطرف الأول المسؤولية عن التأخير أو الأضرار الناجمة عن القوة القاهرة (على سبيل المثال، الكوارث الطبيعية، الإضرابات، الحرائق، عدم توفر قطع الغيار).',
                 'لا تغطي المسؤولية الوحدات غير المدرجة في هذا العقد.',
                 'إذا حالت القوة القاهرة دون التنفيذ لمدة أربعة (4) أشهر أو أكثر، يجوز لأي من الطرفين إنهاء الاتفاقية عن طريق إشعار كتابي.'])
numbered('5', 'Insurance: ', 'The Second Party shall insure its property and equipment against any risk of damage, fire, or accident. The First Party\'s liability is limited to workmanship only.',
         'التأمين: ', 'يلتزم الطرف الثاني بتأمين ممتلكاته ومعداته ضد أي خطر تلف أو حريق أو حادث. تقتصر مسؤولية الطرف الأول على جودة العمل فقط.')
numbered('6', 'Confidentiality: ', 'All contract documents remain the property of the First Party and must not be shared without written approval.',
         'السرية: ', 'تظل جميع مستندات العقد ملكاً للطرف الأول ولا يجوز مشاركتها دون موافقة خطية.')
numbered('7', 'Exclusion of Liabilities: ', 'The First Party, its branches, or its employees shall not be liable for any claim or any direct or indirect damages arising from the Second Party, and it is the responsibility of the Second Party to prove otherwise.',
         'استبعاد المسؤوليات: ', 'لا يتحمل الطرف الأول أو فروعه أو العاملين لديه مسؤولية أي مطالبة أو أي أضرار مباشر أو غير مباشر تنشأ بسبب الطرف الثاني وتقع مسؤولية إثبات عكس ذلك على الطرف الثاني.')
numbered('8', 'Waiver: ', 'Any delay or failure by the First Party to exercise any of its rights under this Contract shall not constitute a waiver or loss of such right or rights. No waiver shall be effective unless it is in writing and signed by the First Party.',
         'التخلي: ', 'لا يشكل أي تأخير أو عدم ممارسة من الطرف الأول لأي حق من حقوقه بموجب هذا العقد تخلياً عن هذا الحق أو الحقوق أو فقدانها ولا يسري أي تخلٍ ما لم يكن كتابياً وموقَّعاً من قِبل الطرف الأول.')
numbered('9', 'Language: ', 'This contract has been prepared in two copies in Arabic and English. In the event of any conflict or interpretation, the Arabic language shall prevail.',
         'اللغة: ', 'تم تحرير هذا العقد من نسختين باللغة العربية والإنجليزية وفي حال التعارض او التفسير تكون اللغة العربية هي الأساس في التفسير.')
numbered('10', 'Calendar: ', 'The Gregorian calendar is adopted in this contract.',
         'التقويم: ', 'التقويم الميلادي هو المعتمد في هذا العقد.')
numbered('11', 'Communications: ', 'All communications between the parties, including notifications, requests, approvals, offers or claims, shall be made via the email address mentioned in the introduction to this contract.',
         'المراسلات: ', 'تتم كل المراسلات بين الطرفين بما في ذلك الإخطارات، أو الطلبات أو الموافقات أو العروض أو المطالبات عبر البريد الإلكتروني المذكور في مقدمة هذا العقد.')

# ── Section 7: Dispute ──
heading('7. Dispute Resolution', '7. تسوية المنازعات')
text_row('In the event of a dispute, the parties must first attempt to resolve the dispute amicably through negotiation. If a solution is not reached within 10 business days, the dispute will be referred to the competent judicial authorities in Jeddah, Kingdom of Saudi Arabia, in accordance with applicable laws.',
         'في حالة وجود نزاع، يجب على الطرفين أولاً محاولة الحل الودي عن طريق التفاوض. وفي حال لم يتم التوصل لحل خلال 10 أيام عمل تحال المنازعات إلى الجهات القضائية المختصة في مدينة جدة بالمملكة العربية السعودية بموجب القوانين المعمول بها.')

insert_page_break()

# ── Section 8: Agreement Acceptance ──
def fn8e(c):
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    make_run(p, '8. Agreement Acceptance', bold=True, size=11, color=NAVY)
    p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.LEFT; para_space(p2, 1, 1)
    make_run(p2, 'By signing below, both Parties acknowledge and agree to the terms and conditions of this Service Contract Agreement', size=11)
def fn8a(c):
    p = c.paragraphs[0]; set_rtl_para(p)
    make_run(p, '8. قبول الاتفاقية', bold=True, size=11, color=NAVY, rtl=True)
    p2 = c.add_paragraph(); set_rtl_para(p2); para_space(p2, 1, 1)
    make_run(p2, 'من خلال التوقيع أدناه، يقر الطرفان ويوافقان على شروط أحكام اتفاقية عقد الخدمة هذه وإستلام كل طرف نسخته الموقعة.', size=11, rtl=True)
add_row(fn8e, fn8a)

close_current_bilingual_table()

# ═══════════════════════════════════════════════════════════════
# Signature Tables
# ═══════════════════════════════════════════════════════════════
def make_sig_table(title_en, title_ar, name_val, contact_val, mobile_val, extra_label_en, extra_label_ar, extra_val):
    _sp = doc.add_paragraph(); para_space(_sp, 4, 0)
    t = doc.add_table(rows=0, cols=4)
    t.allow_autofit = False
    kill_table_borders(t)
    
    tblPr = t._tbl.tblPr
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(int(Cm(PAGE_W).twips)))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    
    # Title
    r = t.add_row()
    c0 = r.cells[0]; c0.merge(r.cells[1]); c2 = r.cells[2]; c2.merge(r.cells[3])
    set_cell_w(c0, Cm(PAGE_W / 2))
    set_cell_w(c2, Cm(PAGE_W / 2))
    clear_cell_borders(c0); clear_cell_borders(c2)
    set_cell_border(c0, left=NAVY_BORDER, top=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)
    set_cell_border(c2, top=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)
    p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    make_run(p, title_en, bold=True, size=11, color=NAVY)
    p = c2.paragraphs[0]; set_rtl_para(p)
    make_run(p, title_ar, bold=True, size=11, color=NAVY, rtl=True)
    # Rows
    def sig_row(le, val, la, h=0):
        rw = t.add_row()
        cl = rw.cells[0]; cv = rw.cells[1]; cv.merge(rw.cells[2]); ca = rw.cells[3]
        set_cell_w(cl, Cm(PAGE_W / 4))
        set_cell_w(cv, Cm(PAGE_W / 2))
        set_cell_w(ca, Cm(PAGE_W / 4))
        clear_cell_borders(cl); clear_cell_borders(cv); clear_cell_borders(ca)
        set_cell_border(cl, left=NAVY_BORDER, bottom=NAVY_BORDER)
        set_cell_border(cv, left=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)
        set_cell_border(ca, right=NAVY_BORDER, bottom=NAVY_BORDER)
        if h:
            rw.height = Cm(h); rw.height_rule = 1
        p = cl.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        make_run(p, le, size=11)
        p = cv.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        make_run(p, val, bold=True, size=11)
        p = ca.paragraphs[0]; set_rtl_para(p)
        make_run(p, la, size=11, rtl=True)
    sig_row('Name', name_val, 'الاسم')
    sig_row('Signature', '', 'التوقيع', 1.0)
    sig_row('Date', '', 'التاريخ')
    sig_row('Stamp', '', 'الختم', 1.0)
    sig_row('Contact Person', contact_val, 'الشخص المسؤول')
    sig_row('Mobile No.', mobile_val, 'رقم الجوال')
    sig_row(extra_label_en, extra_val, extra_label_ar)

make_sig_table('On behalf of Hussein and Al Hassan G. Shaker for Modern Trading Co. Ltd. (First Party)',
               'نيابة عن شركة الاخوان حسين والحسن غازى شاكر للتجارة الحديثة المحدودة (الطرف الأول)',
               'Nouraldeen Riyad Nofal', svc_coord_person, svc_coord_mobile,
               'Unified Number', 'الرقم الموحد', '800 244 0247')
make_sig_table('On behalf of customer (Second Party)',
               'نيابة عن العميل (الطرف الثاني)',
               customer_name, contact_persons, contact_persons_mobile,
               'Additional Information', 'معلومات اضافية', additional_info)

# ═══════════════════════════════════════════════════════════════
# Annex A
# ═══════════════════════════════════════════════════════════════
add_explicit_page_break()

_sp = doc.add_paragraph(); para_space(_sp, 0, 0)
tA = doc.add_table(rows=1, cols=2)
tA.allow_autofit = False
kill_table_borders(tA)
set_col_widths(tA, [Cm(PAGE_W/2), Cm(PAGE_W/2)])
c0 = tA.cell(0, 0)
c1 = tA.cell(0, 1)
clear_cell_borders(c0); clear_cell_borders(c1)
set_cell_border(c0, left=NAVY_BORDER, top=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)
set_cell_border(c1, top=NAVY_BORDER, right=NAVY_BORDER, bottom=NAVY_BORDER)

p = c0.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
make_run(p, 'Annex A \u2013 Scope of Preventive Maintenance', bold=True, size=11, color=NAVY)
p2 = c0.add_paragraph()
p2.paragraph_format.left_indent = Cm(1.5); p2.paragraph_format.first_line_indent = Cm(-0.4)
make_run(p2, u'\u2022  As Mentioned in Quotation', size=11)

p = c1.paragraphs[0]; set_rtl_para(p)
make_run(p, 'الملحق أ - نطاق الصيانة الوقائية', bold=True, size=11, color=NAVY, rtl=True)
p2 = c1.add_paragraph(); set_rtl_para(p2)
p2.paragraph_format.right_indent = Cm(1.5); p2.paragraph_format.first_line_indent = Cm(-0.4)
make_run(p2, u'\u2022  كما هو مذكور في عرض السعر', size=11, rtl=True)

# ═══════════════════════════════════════════════════════════════
# Annex B: Site Details
# ═══════════════════════════════════════════════════════════════
_sp = doc.add_paragraph(); para_space(_sp, 4, 0)
tB = doc.add_table(rows=0, cols=5)
tB.allow_autofit = False
kill_table_borders(tB)

tblPr = tB._tbl.tblPr
tblLayout = OxmlElement('w:tblLayout')
tblLayout.set(qn('w:type'), 'fixed')
tblPr.append(tblLayout)

tblW = OxmlElement('w:tblW')
tblW.set(qn('w:w'), str(int(Cm(PAGE_W).twips)))
tblW.set(qn('w:type'), 'dxa')
tblPr.append(tblW)

col_widths = [Cm(6.0), Cm(2.5), Cm(4.5), Cm(4.4), Cm(2.0)]

# Title
rt = tB.add_row()
ct = rt.cells[0]
for i in range(1,5): ct.merge(rt.cells[i])
set_cell_w(ct, Cm(PAGE_W))
clear_cell_borders(ct); set_cell_border(ct, top=NAVY_GRID_BORDER, bottom=NAVY_GRID_BORDER, left=NAVY_GRID_BORDER, right=NAVY_GRID_BORDER)
p = ct.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
from docx.enum.text import WD_TAB_ALIGNMENT as _WTAB
from docx.shared import Cm as _Cm
p.paragraph_format.tab_stops.add_tab_stop(_Cm(PAGE_W), _WTAB.RIGHT)
make_run(p, 'Annex B \u2013 Site Details', bold=True, size=11, color=NAVY)
make_run(p, '\t', size=11)
make_run(p, 'الملحق ب - تفاصيل الموقع', bold=True, size=11, color=NAVY, rtl=True)

# Headers
rh = tB.add_row()
for i, lbl in enumerate(['Site Address / عنوان الموقع', 'Brand / ماركة', 'Equipment Type / نوع المعدات', 'Contract Type / نوع العقد', 'Quantity / كمية']):
    c = rh.cells[i]
    set_cell_w(c, col_widths[i])
    set_cell_shading(c, HDR_BG_HEX)
    clear_cell_borders(c); set_cell_border(c, top=NAVY_GRID_BORDER, bottom=NAVY_GRID_BORDER, left=NAVY_GRID_BORDER, right=NAVY_GRID_BORDER)
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = make_run(p, lbl, bold=True, size=11)
    r.font.color.rgb = WHITE

# Data
for i, line in enumerate(contract_lines):
    rd = tB.add_row()
    for j, c in enumerate(rd.cells):
        set_cell_w(c, col_widths[j])
        clear_cell_borders(c); set_cell_border(c, top=NAVY_GRID_BORDER, bottom=NAVY_GRID_BORDER, left=NAVY_GRID_BORDER, right=NAVY_GRID_BORDER)
    p = rd.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if i == 0: make_run(p, site_address, bold=True, size=11)
    p = rd.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, line.get('brand_name',''), bold=True, size=11)
    p = rd.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, line.get('product_name',''), bold=True, size=11)
    p = rd.cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    make_run(p, line.get('contract_type',''), bold=True, size=11)
    p = rd.cells[4].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _qty_raw = line.get('qty_ordered', 0)
    try:
        _qty_f = float(_qty_raw)
        _qty_str = str(int(_qty_f)) if _qty_f == int(_qty_f) else f"{_qty_f:,.2f}"
    except (TypeError, ValueError):
        _qty_str = str(_qty_raw)
    make_run(p, _qty_str, bold=True, size=11)

# Merge site address column
if len(contract_lines) > 1:
    first_cell = tB.rows[2].cells[0]  # first data row
    last_cell = tB.rows[1 + len(contract_lines)].cells[0]
    first_cell.merge(last_cell)

doc.save(output_file)
print(f"OK: {output_file}")
